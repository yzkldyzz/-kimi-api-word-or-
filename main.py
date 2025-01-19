from openai import OpenAI
import requests
import docx
import os
import win32com.client
import time
# 打印docx库的路径，确认是否正确加载
print(docx.__file__)

# 将 .doc 文件转换为 .docx 文件
def convert_doc_to_docx(doc_path):
    if not doc_path.endswith('.doc'):
        return doc_path

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    new_path = doc_path + "x"
    doc.SaveAs2(new_path, FileFormat=16)  # 16 表示保存为 .docx 格式
    doc.Close()
    word.Quit()
    return new_path

# 翻译函数
def kimi_translate(query, language, api_key):
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }
    data = {
        'model': 'moonshot-v1-128k',
        'messages': [
            {'role': 'system', 'content': '你是 Kimi，由 Moonshot AI 提供的人工智能助手，你擅长翻译中文和英文的内容。'},
            {'role': 'user', 'content': f'请将以下内容翻译为{language}:\n{query}'}
        ],
        'temperature': 0.3,
    }
    response = requests.post('https://api.moonshot.cn/v1/chat/completions', headers=headers, json=data)
    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    else:
        raise Exception(f'API调用失败，状态码：{response.status_code}')

# 按页翻译Word文档
def translate_word_document_by_page(input_file, output_file, target_language, api_key):
    doc = docx.Document(input_file)
    translated_doc = docx.Document()

    page_breaks = [0]  # 用于记录分页的起始段落索引

    # 标记分页符所在段落
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "":
            page_breaks.append(i + 1)
    page_breaks.append(len(doc.paragraphs))

    # 分页翻译
    for page_start, page_end in zip(page_breaks[:-1], page_breaks[1:]):
        page_text = "\n".join([doc.paragraphs[i].text for i in range(page_start, page_end) if doc.paragraphs[i].text.strip()])
        if page_text:
            print(f"正在翻译第 {page_breaks.index(page_start) + 1} 页...")
            translated_text = kimi_translate(page_text, target_language, api_key)
            for paragraph in translated_text.split("\n"):
                translated_doc.add_paragraph(paragraph)
                time.sleep(0.1)  # 添加延迟处理
        translated_doc.add_paragraph("")  # 添加空行模拟分页

    translated_doc.save(output_file)
    print(f"文档已翻译并保存至：{output_file}")

# 主函数
if __name__ == "__main__":
    print("欢迎使用按页翻译Word文档工具！")
    input_path = input("请输入需要翻译的Word文档路径（例如 input.docx）：")
    input_path = convert_doc_to_docx(input_path)

    output_path = input("请输入翻译后保存的Word文档路径（例如 output.docx）：").strip()
    target_lang = input("请输入目标语言（例如 '英文' 或 '中文'）：").strip()
    api_key = input("请输入您的Moonshot AI API密钥：")  # 请确保安全地处理您的API密钥

    try:
        translate_word_document_by_page(input_path, output_path, target_lang, api_key)
    except Exception as e:
        print(f"翻译过程中出现错误：{e}")
